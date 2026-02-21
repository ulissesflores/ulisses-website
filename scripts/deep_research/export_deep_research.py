#!/usr/bin/env python3
from __future__ import annotations

from pathlib import Path
import re

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer


ROOT = Path(__file__).resolve().parents[2]
ARTICLE_DIR = ROOT / "public" / "deep-research" / "2017-chaos-theory-economics"
MD_PATH = ARTICLE_DIR / "deep-research.md"
DOCX_PATH = ARTICLE_DIR / "deep-research.docx"
PDF_PATH = ARTICLE_DIR / "deep-research.pdf"


def strip_markdown_inline(text: str) -> str:
    text = text.replace("**", "").replace("*", "")
    text = text.replace("`", "")
    text = re.sub(r"\[(.*?)\]\((.*?)\)", r"\1 (\2)", text)
    return text.strip()


def iter_blocks(markdown: str):
    lines = markdown.splitlines()
    paragraph = []

    def flush_paragraph():
        nonlocal paragraph
        if paragraph:
            yield ("p", " ".join(paragraph).strip())
            paragraph = []

    for raw in lines:
        line = raw.rstrip()
        if line.startswith("# "):
            yield from flush_paragraph()
            yield ("h1", strip_markdown_inline(line[2:]))
            continue
        if line.startswith("## "):
            yield from flush_paragraph()
            yield ("h2", strip_markdown_inline(line[3:]))
            continue
        if line.startswith("### "):
            yield from flush_paragraph()
            yield ("h3", strip_markdown_inline(line[4:]))
            continue
        if line.startswith("- "):
            yield from flush_paragraph()
            yield ("li", strip_markdown_inline(line[2:]))
            continue
        if not line.strip():
            yield from flush_paragraph()
            continue
        paragraph.append(strip_markdown_inline(line))

    yield from flush_paragraph()


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr)
    run._r.append(fld_char_end)


def export_docx(markdown: str) -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style.font.size = Pt(12)
    normal_style.paragraph_format.line_spacing = 2

    header = section.header
    header_paragraph = header.paragraphs[0]
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_page_number(header_paragraph)

    for kind, text in iter_blocks(markdown):
        if kind == "h1":
            p = document.add_heading(text, level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            continue
        if kind == "h2":
            p = document.add_heading(text, level=2)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            continue
        if kind == "h3":
            p = document.add_heading(text, level=3)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            continue
        if kind == "li":
            p = document.add_paragraph(style="List Bullet")
            p.add_run(text)
            continue
        p = document.add_paragraph(text)
        p.paragraph_format.line_spacing = 2

    document.save(DOCX_PATH)


def on_page(canvas, _doc):
    canvas.setFont("Times-Roman", 10)
    width, height = A4
    canvas.drawRightString(width - inch, height - 0.75 * inch, str(canvas.getPageNumber()))


def export_pdf(markdown: str) -> None:
    doc = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=A4,
        rightMargin=inch,
        leftMargin=inch,
        topMargin=inch,
        bottomMargin=inch,
    )
    styles = getSampleStyleSheet()
    base = ParagraphStyle(
        "Base",
        parent=styles["Normal"],
        fontName="Times-Roman",
        fontSize=12,
        leading=24,
        spaceAfter=10,
    )
    heading1 = ParagraphStyle(
        "Heading1",
        parent=styles["Heading1"],
        fontName="Times-Bold",
        fontSize=16,
        leading=22,
        spaceBefore=12,
        spaceAfter=8,
    )
    heading2 = ParagraphStyle(
        "Heading2",
        parent=styles["Heading2"],
        fontName="Times-Bold",
        fontSize=14,
        leading=20,
        spaceBefore=10,
        spaceAfter=6,
    )
    bullet = ParagraphStyle(
        "Bullet",
        parent=base,
        leftIndent=14,
        bulletIndent=4,
    )

    flow = []
    for kind, text in iter_blocks(markdown):
        if kind == "h1":
            flow.append(Paragraph(text, heading1))
            continue
        if kind == "h2":
            flow.append(Paragraph(text, heading2))
            continue
        if kind == "h3":
            flow.append(Paragraph(text, heading2))
            continue
        if kind == "li":
            flow.append(Paragraph(f"- {text}", bullet))
            continue
        flow.append(Paragraph(text, base))
        flow.append(Spacer(1, 4))

    doc.build(flow, onFirstPage=on_page, onLaterPages=on_page)


def main() -> None:
    markdown = MD_PATH.read_text(encoding="utf-8")
    export_docx(markdown)
    export_pdf(markdown)


if __name__ == "__main__":
    main()

